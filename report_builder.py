#!/usr/bin/env python3
"""
Weekly Report Builder v1.0
一键周报生成器 — 输入每日工作记录，AI 整理生成专业 DOCX 周报

用法:
  python report_builder.py                        # 交互式输入
  python report_builder.py --file notes.txt       # 从文件读取
  python report_builder.py --demo                 # 生成示例周报
  
免费模型: GLM-4-Flash (智谱, 永久免费)
"""

import json
import sys
import os
import argparse
from datetime import datetime, timedelta

# ===== 依赖安装提示 =====
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("正在安装依赖...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    import requests
except ImportError:
    os.system("pip install requests -q")
    import requests

# ===== 免费 AI 模型配置 =====
FREE_API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
API_KEY = "c9a977acd8044c34b4db930bfa05949f.l8XJrkkloPXRPwAZ"
FREE_MODEL = "glm-4-flash"

# ===== 颜色主题 =====
COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)    # 深蓝
COLOR_ACCENT = RGBColor(0x2B, 0x7A, 0x78)      # 青绿
COLOR_DARK = RGBColor(0x2D, 0x34, 0x36)         # 近黑
COLOR_MEDIUM = RGBColor(0x63, 0x6E, 0x72)       # 灰色
COLOR_LIGHT_BG = RGBColor(0xF0, 0xF4, 0xF8)     # 浅灰蓝

# ===== AI 调用 =====

def call_free_model(prompt, system_prompt=None):
    """调用免费 GLM-4-Flash 模型"""
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})
    
    try:
        resp = requests.post(
            FREE_API_URL,
            headers={
                "Authorization": f"Bearer {API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": FREE_MODEL,
                "messages": messages,
                "max_tokens": 2000,
                "temperature": 0.3
            },
            timeout=60
        )
        if resp.status_code == 200:
            result = resp.json()
            return result["choices"][0]["message"]["content"]
        else:
            print(f"⚠️ API 返回错误: {resp.status_code}")
            return None
    except Exception as e:
        print(f"⚠️ API 调用失败: {e}")
        return None


def organize_notes(raw_text):
    """AI 整理原始工作记录"""
    system = "你是一个专业的周报写作助手。将以下原始工作记录整理成结构化的周报内容，用中文输出。"
    prompt = f"""请将以下工作记录整理成周报格式，包含:
1. 本周工作内容（分条列出，每条一句话）
2. 重点工作/成果（highlight）
3. 遇到的问题和解决方案
4. 下周计划

原始记录：
{raw_text}

请用简洁的专业语言输出。"""
    
    return call_free_model(prompt, system)


# ===== DOCX 生成 =====

def create_report(data, output_path):
    """生成周报 DOCX"""
    doc = Document()
    
    # 样式设置
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    
    # === 标题 ===
    title = doc.add_heading('', level=0)
    run = title.add_run(f'周工作汇报')
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_PRIMARY
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期
    today = datetime.now()
    week_start = today - timedelta(days=today.weekday())
    week_end = week_start + timedelta(days=6)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'{(week_start - timedelta(days=7)).strftime("%Y.%m.%d")} - {week_start.strftime("%Y.%m.%d")}')
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MEDIUM
    
    doc.add_paragraph()  # 空行
    
    # === 基本信息 ===
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    name_cell = info_table.cell(0, 0)
    name_cell.text = f'姓名：____________\n部门：____________'
    date_cell = info_table.cell(0, 1)
    date_cell.text = f'日期：{today.strftime("%Y-%m-%d")}\n状态：□ 完成  □ 进行中'
    
    doc.add_paragraph()
    
    # === 分割线 ===
    def add_section(title_text):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        run = para.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_PRIMARY
        # 下划线
        para.paragraph_format.space_after = Pt(6)
        # 添加底纹装饰
        return para
    
    # 解析 AI 输出
    sections = data.split('\n')
    current_section = None
    section_content = []
    
    for line in sections:
        line = line.strip()
        if not line:
            continue
        
        # 检测是不是章节标题
        if any(keyword in line for keyword in ['本周工作', '重点工作', '遇到问题', '下周计划', '本周工作内容']):
            if current_section and section_content:
                _write_section(doc, current_section, section_content)
            current_section = line
            section_content = []
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '-', '•', '*')) or line[0].isdigit():
            section_content.append(line)
        elif current_section:
            section_content.append(line)
    
    # 最后一节
    if current_section and section_content:
        _write_section(doc, current_section, section_content)
    
    # === 签名区 ===
    doc.add_paragraph()
    doc.add_paragraph()
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.style = 'Table Grid'
    sign_table.cell(0, 0).text = '汇报人签名：____________'
    sign_table.cell(0, 1).text = f'日期：{today.strftime("%Y-%m-%d")}'
    
    doc.save(output_path)
    return output_path


def _write_section(doc, title, items):
    """写入一个章节"""
    # 清理标题
    clean_title = title.strip(' #*-')
    if len(clean_title) > 30:
        clean_title = clean_title[:30]
    
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    run = para.add_run(clean_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_PRIMARY
    
    for item in items:
        item = item.strip()
        if not item:
            continue
        # 清理序号前缀
        clean_item = item
        for prefix in ['1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.', '- ', '• ', '* ']:
            if clean_item.startswith(prefix):
                clean_item = clean_item[len(prefix):]
                break
        
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(clean_item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_DARK


# ===== 主程序 =====

def main():
    parser = argparse.ArgumentParser(description='一键周报生成器 - AI 整理工作记录生成专业 DOCX 周报')
    parser.add_argument('--file', '-f', help='从文件读取工作记录')
    parser.add_argument('--demo', action='store_true', help='生成示例周报')
    parser.add_argument('--output', '-o', default=None, help='输出文件路径')
    
    args = parser.parse_args()
    
    print("""
    ╔═══════════════════════════════════╗
    ║    一键周报生成器 v1.0            ║
    ║    AI 整理 · 免费 · 零成本        ║
    ╚═══════════════════════════════════╝
    """)
    
    raw_notes = None
    
    if args.demo:
        raw_notes = """周一：开会讨论项目进度，修改了登录页面的bug，review了同事的代码
周二：完成了用户管理模块的开发，写了单元测试
周三：和产品经理对接新需求，评估了工作量，开始写设计文档
周四：继续开发搜索功能，解决了分页查询的性能问题（加了索引后速度提升10倍）
周五：部署了新版本到测试环境，发现了一个线上问题并紧急修复"""
        print("📋 使用示例数据...")
    
    elif args.file:
        with open(args.file, 'r', encoding='utf-8') as f:
            raw_notes = f.read()
        print(f"📋 从文件读取: {args.file}")
    
    else:
        print("📝 请输入你的本周工作记录（每行一条，输入空行结束）：")
        lines = []
        while True:
            line = input()
            if not line:
                break
            lines.append(line)
        raw_notes = '\n'.join(lines)
    
    if not raw_notes or not raw_notes.strip():
        print("❌ 没有输入内容")
        return
    
    # AI 整理
    print("🧠 AI 正在整理你的工作记录...（使用免费 GLM-4-Flash 模型）")
    organized = organize_notes(raw_notes)
    
    if not organized:
        print("❌ AI 整理失败，使用原始内容直接生成")
        organized = raw_notes
    
    print("✅ AI 整理完成")
    
    # 生成 DOCX
    output_path = args.output or f'周报_{datetime.now().strftime("%Y%m%d")}.docx'
    print(f"📄 生成周报文档...")
    
    create_report(organized, output_path)
    
    size = os.path.getsize(output_path)
    print(f"✅ 周报已生成: {output_path} ({size/1024:.1f} KB)")
    print(f"💡 打开即可查看和打印")

if __name__ == '__main__':
    main()
